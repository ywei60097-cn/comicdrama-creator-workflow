from io import BytesIO
from base64 import b64encode

from docx import Document
from fastapi.testclient import TestClient

from comicdrama.api.app import app
from comicdrama.core.importers import extract_source_text


def test_extract_txt_source():
    result = extract_source_text("demo.txt", "林舟说：你好".encode("utf-8"))

    assert result.title == "demo"
    assert result.source_format == "txt"
    assert "林舟" in result.text


def test_extract_docx_source():
    document = Document()
    document.add_paragraph("林舟在钟楼下等雨停。")
    document.add_paragraph("苏晴说：我们还有六个小时。")
    buffer = BytesIO()
    document.save(buffer)

    result = extract_source_text("雨夜来信.docx", buffer.getvalue())

    assert result.title == "雨夜来信"
    assert result.source_format == "docx"
    assert "林舟在钟楼下等雨停" in result.text
    assert "苏晴说" in result.text


def test_legacy_doc_has_clear_error():
    try:
        extract_source_text("old.doc", b"not a docx")
    except ValueError as exc:
        assert "Convert it to .docx or PDF" in str(exc)
    else:
        raise AssertionError("legacy .doc should not be accepted")


def test_extract_text_api_accepts_docx():
    document = Document()
    document.add_paragraph("林舟上传了一份文档。")
    buffer = BytesIO()
    document.save(buffer)

    response = TestClient(app).post(
        "/api/v1/files/extract-text",
        json={
            "filename": "项目设定.docx",
            "content_base64": b64encode(buffer.getvalue()).decode("ascii"),
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["source_format"] == "docx"
    assert "林舟上传" in body["text"]
