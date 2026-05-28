from __future__ import annotations

import csv
import json
from dataclasses import dataclass, field
from html.parser import HTMLParser
from io import BytesIO, StringIO
from pathlib import Path
from typing import List, Optional


SUPPORTED_SOURCE_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".pdf",
    ".docx",
    ".html",
    ".htm",
    ".csv",
    ".json",
    ".rtf",
    ".xlsx",
}

UNSUPPORTED_SOURCE_EXTENSIONS = {
    ".doc": "Legacy .doc is a binary Word format. Convert it to .docx or PDF before importing.",
    ".pages": "Apple Pages files should be exported to .docx or PDF before importing.",
}


@dataclass
class ExtractedSource:
    filename: str
    title: str
    source_format: str
    text: str
    pages: Optional[int] = None
    sheets: Optional[int] = None
    notices: List[str] = field(default_factory=list)


def extract_source_text(filename: str, raw: bytes) -> ExtractedSource:
    suffix = Path(filename).suffix.lower()
    title = Path(filename).stem or "Untitled Novel"
    if suffix in {".txt", ".md", ".markdown"}:
        return ExtractedSource(filename, title, suffix.lstrip("."), _decode_text(raw))
    if suffix == ".pdf":
        text, pages = _extract_pdf_text(raw)
        notices = []
        if not text.strip():
            notices.append("No selectable text was found. Scanned PDFs require OCR before processing.")
        return ExtractedSource(filename, title, "pdf", text, pages=pages, notices=notices)
    if suffix == ".docx":
        text = _extract_docx_text(raw)
        return ExtractedSource(filename, title, "docx", text)
    if suffix in {".html", ".htm"}:
        return ExtractedSource(filename, title, suffix.lstrip("."), _extract_html_text(raw))
    if suffix == ".csv":
        return ExtractedSource(filename, title, "csv", _extract_csv_text(raw))
    if suffix == ".json":
        return ExtractedSource(filename, title, "json", _extract_json_text(raw))
    if suffix == ".rtf":
        return ExtractedSource(filename, title, "rtf", _extract_rtf_text(raw))
    if suffix == ".xlsx":
        text, sheets = _extract_xlsx_text(raw)
        return ExtractedSource(filename, title, "xlsx", text, sheets=sheets)
    if suffix in UNSUPPORTED_SOURCE_EXTENSIONS:
        raise ValueError(UNSUPPORTED_SOURCE_EXTENSIONS[suffix])
    supported = ", ".join(sorted(SUPPORTED_SOURCE_EXTENSIONS))
    raise ValueError(f"Unsupported file type: {suffix or 'unknown'}. Supported formats: {supported}.")


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file. Use UTF-8, GB18030, or Big5.")


def _extract_pdf_text(raw: bytes) -> tuple[str, int]:
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires PyPDF2. Install project dependencies first.") from exc

    reader = PdfReader(BytesIO(raw))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"## Page {index}\n\n{page_text.strip()}")
    return "\n\n".join(pages).strip(), len(reader.pages)


def _extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx. Install project dependencies first.") from exc

    document = Document(BytesIO(raw))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


class _TextHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: List[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip_depth += 1
        if tag in {"p", "div", "br", "li", "h1", "h2", "h3", "h4", "tr"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._skip_depth:
            self._skip_depth -= 1
        if tag in {"p", "div", "li", "h1", "h2", "h3", "h4", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self._skip_depth and data.strip():
            self.parts.append(data.strip())


def _extract_html_text(raw: bytes) -> str:
    parser = _TextHTMLParser()
    parser.feed(_decode_text(raw))
    lines = [line.strip() for line in "\n".join(parser.parts).splitlines() if line.strip()]
    return "\n\n".join(lines)


def _extract_csv_text(raw: bytes) -> str:
    text = _decode_text(raw)
    rows = csv.reader(StringIO(text))
    return "\n".join(" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows)


def _extract_json_text(raw: bytes) -> str:
    payload = json.loads(_decode_text(raw))
    lines: List[str] = []
    _flatten_json(payload, lines)
    return "\n".join(lines)


def _flatten_json(value: object, lines: List[str], prefix: str = "") -> None:
    if isinstance(value, dict):
        for key, item in value.items():
            _flatten_json(item, lines, f"{prefix}{key}: " if not prefix else f"{prefix}.{key}: ")
    elif isinstance(value, list):
        for index, item in enumerate(value, start=1):
            _flatten_json(item, lines, f"{prefix}[{index}] ")
    elif value is not None:
        lines.append(f"{prefix}{value}".strip())


def _extract_rtf_text(raw: bytes) -> str:
    text = _decode_text(raw)
    output: List[str] = []
    index = 0
    while index < len(text):
        char = text[index]
        if char in "{}":
            index += 1
            continue
        if char == "\\":
            index += 1
            while index < len(text) and text[index].isalpha():
                index += 1
            while index < len(text) and (text[index].isdigit() or text[index] == "-"):
                index += 1
            if index < len(text) and text[index] == " ":
                index += 1
            continue
        output.append(char)
        index += 1
    return "\n".join(line.strip() for line in "".join(output).splitlines() if line.strip())


def _extract_xlsx_text(raw: bytes) -> tuple[str, int]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("XLSX support requires openpyxl. Install the optional spreadsheet dependency first.") from exc

    workbook = load_workbook(BytesIO(raw), read_only=True, data_only=True)
    sections: List[str] = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            sections.append(f"## Sheet: {sheet.title}\n\n" + "\n".join(rows))
    return "\n\n".join(sections), len(workbook.worksheets)
